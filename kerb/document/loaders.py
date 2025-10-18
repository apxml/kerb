"""Document loaders for various file formats.

This module provides format-specific document loaders for:
- Text files (TXT)
- Markdown files (MD)
- JSON files
- CSV files
- XML files
- HTML files
- PDF files
- DOCX files

Each loader returns a Document object with content and metadata.
"""

import json
import os
import re
from pathlib import Path
from typing import Any, Dict

from kerb.core.types import Document, DocumentFormat


def load_document(file_path: str, **kwargs) -> Document:
    """Load a document from file, automatically detecting format.

    This is the main entry point for loading documents. It detects the format
    and delegates to the appropriate loader.

    Args:
        file_path (str): Path to the document file
        **kwargs: Additional arguments passed to format-specific loaders

    Returns:
        Document: Loaded document with content and metadata

    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If format is not supported

    Examples:
        >>> doc = load_document("report.pdf")
        >>> print(doc.content[:100])

        >>> doc = load_document("data.csv", parse_as_dict=True)
        >>> print(doc.metadata['rows'])
    """
    from .utils import detect_format

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    fmt = detect_format(file_path)

    loaders = {
        DocumentFormat.PDF: load_pdf,
        DocumentFormat.DOCX: load_docx,
        DocumentFormat.HTML: load_html,
        DocumentFormat.MARKDOWN: load_markdown,
        DocumentFormat.TXT: load_text,
        DocumentFormat.CSV: load_csv,
        DocumentFormat.JSON: load_json,
        DocumentFormat.XML: load_xml,
    }

    loader = loaders.get(fmt)
    if loader is None:
        raise ValueError(f"Unsupported format: {fmt.value}")

    doc = loader(file_path, **kwargs)
    doc.format = fmt
    doc.source = file_path

    return doc


def load_text(file_path: str, encoding: str = "utf-8") -> Document:
    """Load a plain text file.

    Args:
        file_path (str): Path to text file
        encoding (str): Text encoding. Defaults to 'utf-8'.

    Returns:
        Document: Loaded document

    Examples:
        >>> doc = load_text("notes.txt")
        >>> print(doc.content)
    """
    with open(file_path, "r", encoding=encoding) as f:
        content = f.read()

    metadata = {
        "encoding": encoding,
        "size": os.path.getsize(file_path),
        "lines": content.count("\n") + 1,
    }

    return Document(content=content, metadata=metadata)


def load_markdown(file_path: str, extract_frontmatter: bool = True) -> Document:
    """Load a Markdown file.

    Args:
        file_path (str): Path to markdown file
        extract_frontmatter (bool): Extract YAML frontmatter if present

    Returns:
        Document: Loaded document with frontmatter in metadata

    Examples:
        >>> doc = load_markdown("README.md")
        >>> if 'frontmatter' in doc.metadata:
        ...     print(doc.metadata['frontmatter'])
    """
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    metadata = {
        "size": os.path.getsize(file_path),
    }

    # Extract frontmatter if present
    if extract_frontmatter:
        frontmatter_pattern = r"^---\s*\n(.*?)\n---\s*\n"
        match = re.match(frontmatter_pattern, content, re.DOTALL)
        if match:
            frontmatter_text = match.group(1)
            metadata["frontmatter"] = {}

            # Simple YAML parsing (for basic key-value pairs)
            for line in frontmatter_text.split("\n"):
                if ":" in line:
                    key, value = line.split(":", 1)
                    metadata["frontmatter"][key.strip()] = value.strip()

            # Remove frontmatter from content
            content = content[match.end() :]

    # Extract headings
    headings = re.findall(r"^#{1,6}\s+(.+)$", content, re.MULTILINE)
    metadata["headings"] = headings

    return Document(content=content, metadata=metadata)


def load_json(file_path: str, as_string: bool = False) -> Document:
    """Load a JSON file.

    Args:
        file_path (str): Path to JSON file
        as_string (bool): If True, return formatted JSON as string content.
                         If False, store parsed object in metadata.

    Returns:
        Document: Loaded document

    Examples:
        >>> doc = load_json("data.json", as_string=True)
        >>> print(doc.content)

        >>> doc = load_json("config.json")
        >>> print(doc.metadata['json_data'])
    """
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    if as_string:
        content = json.dumps(data, indent=2)
    else:
        content = str(data)

    metadata = {
        "json_data": data,
        "size": os.path.getsize(file_path),
    }

    return Document(content=content, metadata=metadata)


def load_csv(
    file_path: str, parse_as_dict: bool = True, encoding: str = "utf-8"
) -> Document:
    """Load a CSV file.

    Args:
        file_path (str): Path to CSV file
        parse_as_dict (bool): Parse CSV and store structured data in metadata
        encoding (str): Text encoding

    Returns:
        Document: Loaded document with CSV data in metadata

    Examples:
        >>> doc = load_csv("data.csv")
        >>> rows = doc.metadata['rows']
        >>> headers = doc.metadata['headers']
    """
    with open(file_path, "r", encoding=encoding) as f:
        content = f.read()

    metadata = {
        "size": os.path.getsize(file_path),
        "encoding": encoding,
    }

    if parse_as_dict:
        lines = content.strip().split("\n")
        if lines:
            headers = [h.strip() for h in lines[0].split(",")]
            rows = []

            for line in lines[1:]:
                values = [v.strip() for v in line.split(",")]
                if len(values) == len(headers):
                    rows.append(dict(zip(headers, values)))

            metadata["headers"] = headers
            metadata["rows"] = rows
            metadata["num_rows"] = len(rows)

    return Document(content=content, metadata=metadata)


def load_xml(file_path: str, encoding: str = "utf-8") -> Document:
    """Load an XML file.

    Args:
        file_path (str): Path to XML file
        encoding (str): Text encoding

    Returns:
        Document: Loaded document

    Examples:
        >>> doc = load_xml("data.xml")
        >>> print(doc.content)
    """
    with open(file_path, "r", encoding=encoding) as f:
        content = f.read()

    # Extract root tag
    root_match = re.search(r"<(\w+)", content)
    root_tag = root_match.group(1) if root_match else None

    metadata = {
        "size": os.path.getsize(file_path),
        "encoding": encoding,
        "root_tag": root_tag,
    }

    return Document(content=content, metadata=metadata)


def load_html(
    file_path: str, extract_text: bool = True, encoding: str = "utf-8"
) -> Document:
    """Load an HTML file.

    Args:
        file_path (str): Path to HTML file
        extract_text (bool): If True, extract plain text from HTML
        encoding (str): Text encoding

    Returns:
        Document: Loaded document

    Examples:
        >>> doc = load_html("page.html", extract_text=True)
        >>> print(doc.content)  # Plain text without HTML tags
    """
    from .extractors import extract_text_from_html

    with open(file_path, "r", encoding=encoding) as f:
        html_content = f.read()

    content = html_content
    metadata = {
        "size": os.path.getsize(file_path),
        "encoding": encoding,
        "raw_html": html_content,
    }

    if extract_text:
        # Basic HTML text extraction
        content = extract_text_from_html(html_content)
        metadata["extracted_text"] = True

    # Extract title
    title_match = re.search(
        r"<title>(.*?)</title>", html_content, re.IGNORECASE | re.DOTALL
    )
    if title_match:
        metadata["title"] = title_match.group(1).strip()

    return Document(content=content, metadata=metadata)


def load_pdf(file_path: str, extract_images: bool = False) -> Document:
    """Load a PDF file.

    Requires: pypdf or PyPDF2 package

    Args:
        file_path (str): Path to PDF file
        extract_images (bool): Whether to extract image information

    Returns:
        Document: Loaded document with page-by-page content

    Examples:
        >>> doc = load_pdf("report.pdf")
        >>> print(f"Pages: {doc.metadata['num_pages']}")
        >>> print(doc.content)  # All pages concatenated
    """
    try:
        import pypdf

        PdfReader = pypdf.PdfReader
    except ImportError:
        try:
            import PyPDF2

            PdfReader = PyPDF2.PdfReader
        except ImportError:
            raise ImportError(
                "PDF support requires pypdf or PyPDF2. "
                "Install with: pip install pypdf"
            )

    reader = PdfReader(file_path)

    pages = []
    for page in reader.pages:
        pages.append(page.extract_text())

    content = "\n\n".join(pages)

    metadata = {
        "num_pages": len(pages),
        "size": os.path.getsize(file_path),
    }

    # Extract PDF metadata
    if reader.metadata:
        pdf_meta = {}
        for key in ["/Title", "/Author", "/Subject", "/Creator", "/Producer"]:
            if key in reader.metadata:
                pdf_meta[key.lstrip("/")] = reader.metadata[key]
        if pdf_meta:
            metadata["pdf_metadata"] = pdf_meta

    return Document(content=content, metadata=metadata, page_content=pages)


def load_docx(file_path: str) -> Document:
    """Load a DOCX file.

    Requires: python-docx package

    Args:
        file_path (str): Path to DOCX file

    Returns:
        Document: Loaded document

    Examples:
        >>> doc = load_docx("report.docx")
        >>> print(doc.content)
    """
    try:
        import docx
    except ImportError:
        raise ImportError(
            "DOCX support requires python-docx. "
            "Install with: pip install python-docx"
        )

    document = docx.Document(file_path)

    # Extract paragraphs
    paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
    content = "\n\n".join(paragraphs)

    metadata = {
        "num_paragraphs": len(paragraphs),
        "size": os.path.getsize(file_path),
    }

    # Extract core properties
    props = document.core_properties
    doc_metadata = {}
    for prop in ["title", "author", "subject", "keywords", "created", "modified"]:
        if hasattr(props, prop):
            value = getattr(props, prop)
            if value:
                doc_metadata[prop] = str(value)

    if doc_metadata:
        metadata["document_properties"] = doc_metadata

    # Extract tables if present
    if document.tables:
        metadata["num_tables"] = len(document.tables)

    return Document(content=content, metadata=metadata)
